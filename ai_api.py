# ai_explainer.py
import os
import pandas as pd
from huggingface_hub import InferenceClient
import pypandoc
import markdown
from docx import Document
from docx.shared import Inches
from binary_utils import get_pandoc_bin

def generate_ai_explanation(patient_folder, api_key=None):
    """Generate AI explanation for the patient's report.csv and save as DOCX."""
    report_path = os.path.join(patient_folder, "report.csv")
    if not os.path.exists(report_path):
        raise FileNotFoundError(f"No report.csv found in {patient_folder}")

    # Load patient report
    df = pd.read_csv(report_path)
    row = df.iloc[0].to_dict()
    report_summary = "\n".join([f"{k}: {v}" for k, v in row.items()])

    # HuggingFace client
    client = InferenceClient(
        provider="cerebras",
        api_key=api_key or os.environ.get("HF_API_TOKEN", "Your_api_key"),
    )

    # Prompt
    prompt = f"""
    You are the Sarcopenia Detection Assistant — an intelligent medical imaging tool that has just analyzed the patient's CT scan.
    You performed the image conversion, segmentation, and measurement of muscle and fat areas, and calculated the Sarcopenia risk.

    Here is the patient's body composition report:

    {report_summary}

    Your task:
    1. Explain the results as if you, the assistant, personally performed the scan analysis and calculations.
    2. Walk the patient step by step through the values (muscle area, fat areas, SMI, etc.) — explaining what each number means in simple, supportive terms.
    3. Clearly state whether the values indicate sarcopenia or not, and explain why based on thresholds.
    4. Speak in a caring, professional, and confident tone — reassuring the patient that the analysis was performed accurately.
    5. End with encouragement, practical lifestyle advice (exercise, diet, follow-up with doctor), and emphasize that these results are an aid for discussion with their healthcare provider.
    """


    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
    )

    ai_text = completion.choices[0].message["content"]

    html_text = markdown.markdown(ai_text, extensions=["tables", "fenced_code"])

    #print(html_text)

    # Save to DOCX
    output_path = os.path.join(patient_folder, "final_output.docx")

    # Convert Markdown (from AI) to DOCX
    pypandoc.pandoc_path = get_pandoc_bin()
    pypandoc.convert_text(
        html_text,                     # input text (Markdown with **bold**)
        'docx',                      # output format
        format='html',                 # input format is Markdown
        outputfile=output_path, 
        extra_args=['--standalone']  # ensures complete doc structure
    )
    # Add overlays
    overlay_dir = os.path.join(patient_folder, "overlays")
    if os.path.exists(overlay_dir):
        add_overlays_to_doc(output_path, overlay_dir)

    return output_path

def add_overlays_to_doc(docx_path, overlay_dir):
    doc = Document(docx_path)

    # Add a heading before images
    doc.add_heading("Segmentation Overlays", level=1)

    # Insert each overlay PNG
    for img_file in sorted(os.listdir(overlay_dir)):
        if img_file.endswith(".png"):
            doc.add_picture(os.path.join(overlay_dir, img_file), width=Inches(4))
            doc.add_paragraph(img_file)  # caption

    # Save back
    doc.save(docx_path)
    print(f"✅ Overlays added to {docx_path}")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generate AI explanation report")
    parser.add_argument("--patient_folder", required=True, help="Path to patient folder")
    # parser.add_argument("--api_key", required=False, help="Hugging Face API key")
    args = parser.parse_args()

    explanation = generate_ai_explanation(args.patient_folder)
    print("\n✅ AI Explanation Generated:\n")
    print(explanation)
